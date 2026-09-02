"""
postprocess_docx.py — Generalized DOCX post-processor for pdf2docx output.

Applies three targeted fixes:
1. FILLED-RECT CALLOUT BOXES  — detects filled colored rectangles in the source PDF
   that contain text, and reproduces them as OOXML shaded paragraph blocks in the
   DOCX at the correct position.

2. TABLE COLUMN PROPORTIONS   — for each table in the DOCX, derives the true column
   widths from the source PDF's text-span X-positions (which are layout-exact) and
   rewrites the DOCX tblGrid element with the correct proportional widths.

3. HEADING UNDERLINES         — detects text spans in the PDF that are underlined
   (by checking for underline drawing paths directly beneath them) and adds OOXML
   <w:u w:val="single"/> to matching paragraph runs in the DOCX.

All three fixes are fully generalized:
- No hard-coded coordinates, colors, or document-specific logic.
- Column widths come from measured PDF geometry, not assumptions.
- Callout detection uses a threshold on fill-color saturation + rectangle area.
- Underline detection uses spatial proximity between text baselines and drawing paths.

Usage:
    python postprocess_docx.py <input.pdf> <input.docx> <output.docx>
"""

import sys, os, io, copy, zipfile, json, re, shutil
from pathlib import Path
from collections import defaultdict

import pymupdf                          # PyMuPDF
from lxml import etree
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ══════════════════════════════════════════════════════════════════════════════
# Constants / helpers
# ══════════════════════════════════════════════════════════════════════════════

TWIP = 20          # 1 pt = 20 twips
EMU  = 12700       # 1 pt = 12700 EMU


def rgb_float_to_hex(rgb):
    """(0.573, 0.816, 0.314) → '92d050'"""
    return ''.join(f'{min(255,int(round(c*255))):02x}' for c in rgb)


def cluster_coords(coords, tolerance=3.0):
    if not coords:
        return []
    coords = sorted(set(round(c, 3) for c in coords))
    clusters = [[coords[0]]]
    for c in coords[1:]:
        if c - clusters[-1][-1] <= tolerance:
            clusters[-1].append(c)
        else:
            clusters.append([c])
    return [round(sum(cl) / len(cl), 3) for cl in clusters]


def color_saturation(rgb):
    """Return 0..1 saturation of an RGB tuple (0..1 each)."""
    r, g, b = rgb
    mx, mn = max(r, g, b), min(r, g, b)
    return (mx - mn) / mx if mx > 0 else 0


def is_colored_fill(rgb, min_saturation=0.20, min_brightness=0.05):
    """True if this fill color is distinctly non-white and non-black/grey."""
    if rgb is None:
        return False
    r, g, b = rgb
    brightness = max(r, g, b)
    if brightness < min_brightness:   # too dark (near-black)
        return False
    if all(abs(c - r) < 0.08 for c in [g, b]):  # near-grey (equal channels)
        return False
    return color_saturation(rgb) >= min_saturation


# ══════════════════════════════════════════════════════════════════════════════
# FIX 1 — Filled callout box detection & OOXML reconstruction
# ══════════════════════════════════════════════════════════════════════════════

def detect_callout_boxes(pdf_doc):
    """
    Scan all pages for filled colored rectangles that contain text.
    Returns list of callout box dicts, ordered by (page, y-position).
    Each dict:
        page_no       : 0-based page index
        rect          : (x0, y0, x1, y1) in pt
        fill_hex      : e.g. '007f00'
        text_runs     : list of {text, font, size_pt, bold, color_int}
        full_text     : joined text
        has_border    : bool
        border_hex    : hex or None
        align         : 'center'|'left'|'right'  (estimated from text position)
    """
    boxes = []
    for pno in range(pdf_doc.page_count):
        page = pdf_doc[pno]
        drawings = page.get_drawings()
        
        for d in drawings:
            rect = d.get("rect")
            fill = d.get("fill")
            if rect is None or fill is None:
                continue
            if not is_colored_fill(fill):
                continue
            
            w = rect[2] - rect[0]
            h = rect[3] - rect[1]
            
            # Must be a meaningful rectangle (not a tiny icon or line)
            if w < 20 or h < 8:
                continue
            
            # Extract text that overlaps this rect
            clip = pymupdf.Rect(rect)
            blocks = page.get_text("dict", clip=clip, flags=0)["blocks"]
            text_runs = []
            for b in blocks:
                if b.get("type") != 0:
                    continue
                for line in b.get("lines", []):
                    for span in line.get("spans", []):
                        txt = span.get("text", "").strip()
                        if not txt:
                            continue
                        text_runs.append({
                            "text": txt,
                            "font": span.get("font", ""),
                            "size_pt": round(span.get("size", 11), 1),
                            "color_int": span.get("color", 0),
                            "bbox": span.get("bbox", []),
                        })
            
            # Only keep boxes that actually contain text
            if not text_runs:
                continue
            
            full_text = " ".join(r["text"] for r in text_runs)
            
            # Estimate text alignment within box
            text_x_centers = []
            for r in text_runs:
                bb = r.get("bbox", [])
                if len(bb) >= 4:
                    text_x_centers.append((bb[0] + bb[2]) / 2)
            box_cx = (rect[0] + rect[2]) / 2
            if text_x_centers:
                avg_text_cx = sum(text_x_centers) / len(text_x_centers)
                rel = (avg_text_cx - rect[0]) / (rect[2] - rect[0])
                if rel < 0.35:
                    align = "left"
                elif rel > 0.65:
                    align = "right"
                else:
                    align = "center"
            else:
                align = "center"
            
            # Border
            stroke = d.get("color")
            has_border = (stroke is not None and not all(abs(c - fill[i]) < 0.05 for i, c in enumerate(stroke)))
            
            boxes.append({
                "page_no": pno,
                "rect": (round(rect[0],2), round(rect[1],2), round(rect[2],2), round(rect[3],2)),
                "fill_hex": rgb_float_to_hex(fill),
                "fill_rgb": [round(v,4) for v in fill],
                "text_runs": text_runs,
                "full_text": full_text,
                "has_border": has_border,
                "border_hex": rgb_float_to_hex(stroke) if has_border and stroke else None,
                "align": align,
                "width_pt": round(w, 2),
                "height_pt": round(h, 2),
            })
    
    return boxes


def make_shaded_paragraph_xml(text, fill_hex, font_name, font_size_pt,
                               bold=False, text_color_hex="FFFFFF",
                               align="center", border_hex=None,
                               width_twips=None, space_before_pt=6, space_after_pt=6):
    """
    Build a <w:p> element for a shaded callout box paragraph.
    Uses paragraph shading + optional border.
    """
    # Build namespace map
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    
    p = OxmlElement('w:p')
    
    # Paragraph properties
    pPr = OxmlElement('w:pPr')
    
    # Alignment
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), align)
    pPr.append(jc)
    
    # Spacing
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(int(space_before_pt * TWIP)))
    spacing.set(qn('w:after'), str(int(space_after_pt * TWIP)))
    pPr.append(spacing)
    
    # Paragraph shading (the background fill)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex.upper())
    pPr.append(shd)
    
    # Paragraph border (optional)
    if border_hex:
        pBdr = OxmlElement('w:pBdr')
        for side in ['top', 'left', 'bottom', 'right']:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), '4')
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), border_hex.upper())
            pBdr.append(el)
        pPr.append(pBdr)
    
    # Indentation (simulate box width via indents from page margins if needed)
    if width_twips:
        ind = OxmlElement('w:ind')
        # We'll set symmetric indents to center the box visually
        # This is approximate — exact box width matching requires frames
        pPr.append(ind)
    
    p.append(pPr)
    
    # Run
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Font
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rPr.append(rFonts)
    
    # Size
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(font_size_pt * 2)))
    rPr.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), str(int(font_size_pt * 2)))
    rPr.append(szCs)
    
    # Bold
    if bold:
        rPr.append(OxmlElement('w:b'))
        rPr.append(OxmlElement('w:bCs'))
    
    # Text color
    color_el = OxmlElement('w:color')
    color_el.set(qn('w:val'), text_color_hex.upper())
    rPr.append(color_el)
    
    r.append(rPr)
    
    t = OxmlElement('w:t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = text
    r.append(t)
    p.append(r)
    
    return p


def inject_callout_boxes(docx_path, pdf_doc, out_path):
    """
    For each callout box detected in the PDF, find the matching paragraph in
    the DOCX (by text content) and replace it with a shaded paragraph.
    Also inserts shaded paragraphs where no match exists but the text appears
    as a first/only paragraph in its section.
    """
    boxes = detect_callout_boxes(pdf_doc)
    if not boxes:
        print("  [callout] No colored callout boxes detected in PDF.")
        shutil.copy2(docx_path, out_path)
        return 0

    # Deduplicate by text content — keep largest box for each unique text+fill combo
    seen_keys = {}
    for b in boxes:
        key = (b["full_text"].strip().lower()[:60], b["fill_hex"])
        if key not in seen_keys or b["width_pt"] * b["height_pt"] > seen_keys[key]["width_pt"] * seen_keys[key]["height_pt"]:
            seen_keys[key] = b
    boxes = list(seen_keys.values())

    print(f"  [callout] {len(boxes)} unique callout box(es) after deduplication:")
    for b in boxes:
        print(f"    fill=#{b['fill_hex']}  text='{b['full_text'][:60]}'  {b['width_pt']:.0f}x{b['height_pt']:.0f}pt  align={b['align']}")

    doc = Document(docx_path)
    body = doc.element.body
    paragraphs = body.findall('.//' + qn('w:p'))
    
    patches_applied = 0
    
    for box in boxes:
        target_text = box["full_text"].strip()
        if not target_text:
            continue
        
        # Find paragraph(s) whose text matches the callout content
        matched_para = None
        for para in paragraphs:
            para_text = ''.join(
                t.text or '' for t in para.findall('.//' + qn('w:t'))
            ).strip()
            
            # Fuzzy match: the para text should be contained in or close to box text
            if (para_text and (
                para_text in target_text or
                target_text in para_text or
                (len(para_text) > 5 and para_text[:20] in target_text)
            )):
                matched_para = para
                break
        
        if matched_para is None:
            print(f"  [callout] WARNING: no paragraph match for '{target_text[:40]}' — skipping")
            continue
        
        # Skip if the matched paragraph is inside a table cell
        # (table cell shading already handles that case via pdf2docx)
        W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        ancestor = matched_para.getparent()
        in_table = False
        while ancestor is not None:
            if ancestor.tag == f'{{{W_NS}}}tc':
                in_table = True
                break
            ancestor = ancestor.getparent()
        if in_table:
            print(f"  [callout] Skipping '{target_text[:40]}' — already in table cell with shading")
            continue
        
        # Get font info from first text run
        first_run_info = box["text_runs"][0] if box["text_runs"] else {}
        font_name = "Arial"  # fallback
        raw_font = first_run_info.get("font", "")
        if raw_font:
            # Strip Bold/Italic suffixes
            font_name = re.sub(r'[,\-]?(Bold|Italic|Regular|Condensed|Light)\w*', '',
                               raw_font, flags=re.IGNORECASE).strip() or font_name
        font_size = first_run_info.get("size_pt", 11.0)
        
        # Text color: invert from fill if light-on-dark, else keep original
        fill_brightness = max(box["fill_rgb"])
        color_int = first_run_info.get("color_int", 0)
        r_c = (color_int >> 16) & 0xFF
        g_c = (color_int >> 8)  & 0xFF
        b_c =  color_int        & 0xFF
        text_color_hex = f'{r_c:02x}{g_c:02x}{b_c:02x}'
        
        # Build shaded paragraph
        new_p = make_shaded_paragraph_xml(
            text       = target_text,
            fill_hex   = box["fill_hex"],
            font_name  = font_name,
            font_size_pt = font_size,
            bold       = "Bold" in first_run_info.get("font", ""),
            text_color_hex = text_color_hex,
            align      = box["align"],
            border_hex = box["border_hex"],
            space_before_pt = 4,
            space_after_pt  = 4,
        )
        
        # Replace matched paragraph with shaded version
        parent = matched_para.getparent()
        if parent is None:
            print(f"  [callout] WARNING: cannot find parent for '{target_text[:40]}' — skipping")
            continue
        
        # Build index by iterating parent children
        children = list(parent)
        try:
            idx = children.index(matched_para)
        except ValueError:
            print(f"  [callout] WARNING: paragraph not in parent children list — skipping")
            continue
        
        parent.remove(matched_para)
        parent.insert(idx, new_p)
        patches_applied += 1
        print(f"  [callout] Patched: '{target_text[:50]}' → fill=#{box['fill_hex']}")
    
    doc.save(out_path)
    print(f"  [callout] {patches_applied} callout box(es) patched.")
    return patches_applied


# ══════════════════════════════════════════════════════════════════════════════
# FIX 2 — Table column proportions from PDF geometry
# ══════════════════════════════════════════════════════════════════════════════

def extract_pdf_table_col_widths(pdf_doc):
    """
    Extract table column widths from PDF drawing paths.
    
    Strategy (generalized for any bordered table):
    1. Collect all filled/stroked rectangles from the page.
    2. Group them into Y-bands (horizontal rows) by clustering Y-top coordinates.
    3. For each Y-band with ≥3 rectangles, extract unique X-edges (left/right of each cell rect).
    4. Filter out sub-pixel gaps (border thickness) to get true column edges.
    5. Compute column widths and fractions.
    6. Cluster Y-bands with the same column count into logical table groups.
    
    Falls back to text-span X-origin clustering for borderless tables.
    Returns list ordered by (page_no, y_top).
    """
    all_table_regions = []

    for pno in range(pdf_doc.page_count):
        page = pdf_doc[pno]

        # ── Method A: drawing-path borders ────────────────────────────────
        drawings = page.get_drawings()
        all_rects = []
        for d in drawings:
            rect = d.get("rect")
            if rect is None:
                continue
            x0, y0, x1, y1 = rect
            if x1 - x0 < 5 or y1 - y0 < 0.5:
                continue
            all_rects.append((round(x0, 2), round(y0, 2), round(x1, 2), round(y1, 2)))

        # Group rects by Y-top (within 2pt)
        y_groups = {}
        for r in sorted(all_rects, key=lambda x: x[1]):
            placed = False
            for key in list(y_groups.keys()):
                if abs(r[1] - key) <= 2.0:
                    y_groups[key].append(r)
                    placed = True
                    break
            if not placed:
                y_groups[r[1]] = [r]

        # For each Y-band, extract column widths
        drawing_rows = []  # (y_top, n_cols, col_widths, fracs)
        for yk in sorted(y_groups.keys()):
            rects = y_groups[yk]
            if len(rects) < 3:
                continue
            # Collect all X edges
            all_x = sorted(set(round(r[0], 0) for r in rects) |
                           set(round(r[2], 0) for r in rects))
            # Filter: keep edges with gap >= 5pt (cell walls, not sub-pixel borders)
            filtered_x = [all_x[0]]
            for x in all_x[1:]:
                if x - filtered_x[-1] >= 5:
                    filtered_x.append(x)
            if len(filtered_x) < 3:  # need at least 2 columns
                continue
            col_widths = [round(filtered_x[i+1] - filtered_x[i], 2)
                          for i in range(len(filtered_x) - 1)]
            total = sum(col_widths)
            if total < 10:
                continue
            fracs = [round(w / total, 4) for w in col_widths]
            drawing_rows.append({
                "y_top": yk,
                "page_no": pno,
                "x0": filtered_x[0],
                "x1": filtered_x[-1],
                "col_widths": col_widths,
                "fracs": fracs,
                "n_cols": len(col_widths),
            })

        # Group consecutive drawing rows with same n_cols into table regions
        if drawing_rows:
            table_groups = []
            current = [drawing_rows[0]]
            for row in drawing_rows[1:]:
                prev = current[-1]
                # Same table if: same col count, y gap < 50pt, same x bounds (within 5pt)
                if (row["n_cols"] == prev["n_cols"]
                        and row["y_top"] - prev["y_top"] < 50
                        and abs(row["x0"] - prev["x0"]) < 5
                        and abs(row["x1"] - prev["x1"]) < 5):
                    current.append(row)
                else:
                    table_groups.append(current)
                    current = [row]
            table_groups.append(current)

            for grp in table_groups:
                if not grp:
                    continue
                # Average fracs across all rows (should be identical for same table)
                n = grp[0]["n_cols"]
                avg_fracs = [
                    round(sum(r["fracs"][i] for r in grp) / len(grp), 4)
                    for i in range(n)
                ]
                avg_widths = [round(sum(r["col_widths"][i] for r in grp) / len(grp), 2)
                              for i in range(n)]
                total_w = grp[0]["x1"] - grp[0]["x0"]
                y0 = grp[0]["y_top"]
                y1 = grp[-1]["y_top"] + 15  # approx

                all_table_regions.append({
                    "page_no": pno,
                    "method": "drawings",
                    "bbox": (round(grp[0]["x0"], 2), round(y0, 2),
                             round(grp[0]["x1"], 2), round(y1, 2)),
                    "col_widths_pt": avg_widths,
                    "col_fractions": avg_fracs,
                    "n_cols": n,
                    "n_rows": len(grp),
                })

        # ── Method B (fallback): text X-origin clustering ─────────────────
        # Used for tables with no visible borders (e.g. commission bullet section)
        if not drawing_rows:
            text_dict = page.get_text("dict", flags=0)
            row_spans = defaultdict(list)
            for b in text_dict["blocks"]:
                if b.get("type") != 0:
                    continue
                for line in b.get("lines", []):
                    y_mid = round((line["bbox"][1] + line["bbox"][3]) / 2, 1)
                    for span in line.get("spans", []):
                        row_spans[y_mid].append({
                            "x0": span["bbox"][0],
                            "x1": span["bbox"][2],
                            "text": span.get("text", "").strip(),
                        })
            # Find multi-column rows
            x_starts_all = []
            x0_all, x1_all, y_all = [], [], []
            for y_mid, spans in row_spans.items():
                xs = sorted(set(round(s["x0"], 1) for s in spans if s["text"]))
                if len(xs) >= 2:
                    x_starts_all.extend(xs)
                    for s in spans:
                        x0_all.append(s["x0"])
                        x1_all.append(s["x1"])
                        y_all.extend([y_mid - 5, y_mid + 5])
            if x_starts_all:
                col_edges = cluster_coords(x_starts_all, tolerance=5.0)
                filtered_edges = []
                for e in col_edges:
                    if not filtered_edges or e - filtered_edges[-1] >= 8:
                        filtered_edges.append(e)
                if len(filtered_edges) >= 2:
                    tbl_x0 = min(x0_all) - 2
                    tbl_x1 = max(x1_all) + 2
                    col_widths = [round(filtered_edges[i+1] - filtered_edges[i], 2)
                                  for i in range(len(filtered_edges) - 1)]
                    last_w = round(tbl_x1 - filtered_edges[-1], 2)
                    if last_w > 5:
                        col_widths.append(last_w)
                    total_w = sum(col_widths)
                    fracs = [round(w / total_w, 4) for w in col_widths]
                    all_table_regions.append({
                        "page_no": pno,
                        "method": "text",
                        "bbox": (round(tbl_x0, 2), round(min(y_all), 2),
                                 round(tbl_x1, 2), round(max(y_all), 2)),
                        "col_widths_pt": col_widths,
                        "col_fractions": fracs,
                        "n_cols": len(col_widths),
                        "n_rows": len([y for y, s in row_spans.items()
                                       if len([sp for sp in s if sp["text"]]) >= 2]),
                    })

    return all_table_regions



def get_docx_tables(docx_path):
    """
    Read DOCX and return list of (table_element, col_count, current_col_widths_twips, table_width_twips).
    """
    with zipfile.ZipFile(docx_path) as z:
        doc_xml = z.read("word/document.xml")
    tree = etree.fromstring(doc_xml)
    body = tree.find(f'.//{{{qn("w:body")[2:]}}}body') if False else tree.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}body')
    
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    tables_info = []
    for tbl in body.findall(f'{{{W}}}tbl') + body.findall(f'.//{{{W}}}tbl'):
        grid = tbl.find(f'{{{W}}}tblGrid')
        col_w = []
        if grid is not None:
            for gc in grid.findall(f'{{{W}}}gridCol'):
                w = gc.get(f'{{{W}}}w')
                if w:
                    col_w.append(int(w))
        
        tbl_pr = tbl.find(f'{{{W}}}tblPr')
        tbl_w_twips = None
        if tbl_pr is not None:
            tw_el = tbl_pr.find(f'{{{W}}}tblW')
            if tw_el is not None:
                tw = tw_el.get(f'{{{W}}}w')
                if tw:
                    tbl_w_twips = int(tw)
        
        tables_info.append((tbl, len(col_w), col_w, tbl_w_twips))
    return tables_info


def fix_header_title_indent(docx_path, out_path):
    """
    Root cause (found from a direct report, with screenshots, that the
    heading "doesn't look like the original at all"): the header table's
    title-text cell (sharing a row with the district logo image) renders as
    six short, oddly-broken lines --

        Dist. Govt. KP-
        Provincial District
        Accounts Office
        Peshawar Dist.
        Monthly Salary
        Statement (July-2026)

    -- even though the cell itself is a reasonable width (5526 twips,
    ~276pt) and the source PDF wraps the same text into just three clean
    lines at roughly the same width. Measured directly: pdf2docx wrote the
    title as a SINGLE <w:p> with a <w:ind w:left="1288" w:right="1440"
    w:firstLine="378"/> -- 2728 twips of left+right indent, reproducing
    the text's original horizontal offset from the page's left edge in the
    source PDF (where it made sense: the surrounding content was a normal
    single-column page). Once pdf2docx also puts this same text inside a
    TABLE CELL -- whose horizontal position is *already* fully determined
    by the table's own <w:tblInd> and the cell's column boundary -- that
    paragraph indent is purely redundant, and it eats 2728 of the cell's
    5526 twips before a single character is drawn, leaving only ~140pt to
    wrap into instead of the cell's real ~276pt. That's what produced the
    ugly six-line break; it has nothing to do with the cell's own width,
    which was already reasonable.

    Confirmed by testing: zeroing this one paragraph's ind (left/right/
    firstLine) on the real file, with no other change, turned the six-line
    wrap into two clean lines ("Dist. Govt. KP-Provincial District
    Accounts Office" / "Peshawar Dist. Monthly Salary Statement
    (July-2026)") -- close to the source PDF's own three-line break, and a
    real, natural word-wrap at the cell's actual width rather than an
    artificially strangled one.

    Scope, kept deliberately narrow: only paragraphs in a NON-image cell
    that shares its table with an image cell (i.e. exactly the
    header/logo table pattern this was diagnosed on). Every other cell in
    every other table already renders correctly and is left untouched --
    this is not a general "remove all paragraph indents" fix, because for
    an ordinary text cell the indent can be small and legitimate (already
    accounted for as real width-consumption by `narrow_cols` and
    `merged_row_wraps`). Only the specific pattern that caused the
    reported symptom -- a large, page-position-derived indent stranded
    inside a table cell that already provides its own positioning -- is
    touched.
    """
    doc = Document(docx_path)
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    def q(tag):
        return f'{{{W}}}{tag}'

    body = doc.element.body
    fixed_paras = 0

    for tbl in body.findall(q('tbl')):
        has_image_cell = tbl.find(f'.//{{{W}}}drawing') is not None
        if not has_image_cell:
            continue

        for tr in tbl.findall(q('tr')):
            for tc in tr.findall(q('tc')):
                if tc.find(f'.//{{{W}}}drawing') is not None:
                    continue  # this is the image cell itself -- leave it
                for p in tc.findall(q('p')):
                    pPr = p.find(q('pPr'))
                    if pPr is None:
                        continue
                    ind = pPr.find(q('ind'))
                    if ind is None:
                        continue
                    left = float(ind.get(q('left'), '0') or 0)
                    right = float(ind.get(q('right'), '0') or 0)
                    if left <= 0 and right <= 0:
                        continue
                    ind.set(q('left'), '0')
                    ind.set(q('right'), '0')
                    if ind.get(q('firstLine')) is not None:
                        ind.set(q('firstLine'), '0')
                    fixed_paras += 1

    doc.save(out_path)
    print(f"  [title_indent] {fixed_paras} paragraph(s) in image-table text cells "
          f"had a redundant page-position indent cleared, restoring their full cell width for wrapping.")
    return fixed_paras


def patch_column_widths(docx_path, pdf_doc, out_path):
    """
    Rewrite DOCX table column widths using PDF-geometry-derived proportions.
    
    Algorithm:
    1. Extract table regions from PDF (ordered by page, then Y position).
    2. Match each DOCX table to a PDF region by index (reading order).
    3. For matching regions with the same column count, apply the PDF proportions.
    4. For non-matching regions, keep existing widths (safe fallback).
    """
    pdf_regions = extract_pdf_table_col_widths(pdf_doc)
    
    print(f"  [cols] PDF table regions detected: {len(pdf_regions)}")
    for i, r in enumerate(pdf_regions):
        print(f"    PDF[{i}] p{r['page_no']+1}  ncols={r['n_cols']}  widths={r['col_widths_pt']}  fracs={r['col_fractions']}")
    
    doc = Document(docx_path)
    body = doc.element.body
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    
    # Find all top-level tables (not nested)
    docx_tables = body.findall(f'{{{W}}}tbl')
    print(f"  [cols] DOCX top-level tables: {len(docx_tables)}")
    
    patches = 0
    
    # Build column-count index: for each n_cols value, list of PDF regions with that count
    from collections import defaultdict as _dd
    pdf_by_ncols = _dd(list)
    for r in pdf_regions:
        pdf_by_ncols[r['n_cols']].append(r)
    # Usage pointers per n_cols (round-robin through matching regions)
    pdf_usage = {nc: 0 for nc in pdf_by_ncols}
    
    for ti, tbl in enumerate(docx_tables):
        # Get current grid
        grid = tbl.find(f'{{{W}}}tblGrid')
        if grid is None:
            continue
        current_cols = grid.findall(f'{{{W}}}gridCol')
        n_cols = len(current_cols)

        if n_cols == 0:
            continue

        # Skip any table that contains an image (w:drawing) anywhere inside
        # it. This function's whole method -- deriving column widths from
        # PDF *text*-span x-positions -- doesn't apply to a cell holding a
        # picture instead of clustered text, and the "nearest column-count"
        # fallback match (a couple of lines below) can match a table that
        # has nothing to do with the region it's being resized against --
        # confirmed on a real file: a 2-column header row (title text next
        # to a logo image) had no ncols=2 PDF region to match, fell back to
        # the nearest available one (an unrelated ncols=4 wage-data region),
        # and the resulting nonsense width squeezed the logo's cell enough
        # that the (inline, right-aligned) logo rendered overflowing past
        # the page's right edge -- confirmed by rendering and measuring the
        # image's bounding box against the page size. Leaving an
        # image-bearing table's widths exactly as pdf2docx's own base
        # conversion set them avoids that whole failure mode.
        if tbl.find(f'.//{{{W}}}drawing') is not None:
            print(f"  [cols] Table {ti}: contains an image — skipping (width patch doesn't apply to pictures)")
            continue

        # Current total width from tblPr
        tbl_pr = tbl.find(f'{{{W}}}tblPr')
        tbl_w_twips = None
        if tbl_pr is not None:
            tw_el = tbl_pr.find(f'{{{W}}}tblW')
            if tw_el is not None:
                tw = tw_el.get(f'{{{W}}}w')
                if tw and tw != '0':
                    tbl_w_twips = int(tw)
        
        # Current col widths
        current_widths = []
        for gc in current_cols:
            w = gc.get(f'{{{W}}}w')
            current_widths.append(int(w) if w else 0)
        current_total = sum(current_widths)
        
        # Find matching PDF region by column count (exact match first, then nearest)
        matched_region = None
        
        # Exact match
        if n_cols in pdf_by_ncols:
            idx = pdf_usage[n_cols]
            if idx < len(pdf_by_ncols[n_cols]):
                matched_region = pdf_by_ncols[n_cols][idx]
                pdf_usage[n_cols] = idx + 1
        
        # Nearest match (within 1 col)
        if matched_region is None:
            for delta in [1, 2]:
                for nc in [n_cols - delta, n_cols + delta]:
                    if nc in pdf_by_ncols:
                        idx = pdf_usage.get(nc, 0)
                        if idx < len(pdf_by_ncols[nc]):
                            matched_region = pdf_by_ncols[nc][idx]
                            pdf_usage[nc] = idx + 1
                            break
                if matched_region:
                    break
        
        if matched_region is None:
            print(f"  [cols] Table {ti}: no PDF match (ncols={n_cols}) — keeping existing widths")
            continue
        
        # Use the total twips from current DOCX table (or estimate from page width)
        target_total = tbl_w_twips or current_total or (int(5040))  # 5040 twips = ~252mm text width
        
        # Apply PDF proportions to derive new widths
        fracs = matched_region['col_fractions']
        
        # If col counts differ by 1, redistribute
        if len(fracs) != n_cols:
            if len(fracs) > n_cols:
                fracs = fracs[:n_cols]
            else:
                # Pad last column
                remainder = 1.0 - sum(fracs)
                fracs = fracs + [remainder]
            total_frac = sum(fracs)
            fracs = [f / total_frac for f in fracs]
        
        new_widths_twips = []
        for i, frac in enumerate(fracs):
            w = int(round(frac * target_total))
            new_widths_twips.append(max(w, 200))  # minimum 10pt per column
        
        # Adjust last column to hit target exactly
        diff = target_total - sum(new_widths_twips)
        new_widths_twips[-1] += diff
        
        # Update tblGrid
        for i, gc in enumerate(current_cols[:len(new_widths_twips)]):
            gc.set(f'{{{W}}}w', str(new_widths_twips[i]))
        
        # Update each cell's w:tcW in every row
        rows = tbl.findall(f'{{{W}}}tr')
        for row in rows:
            cells = row.findall(f'{{{W}}}tc')
            col_idx = 0
            for cell in cells:
                tc_pr = cell.find(f'{{{W}}}tcPr')
                if tc_pr is None:
                    continue
                
                # Handle gridSpan — how many grid columns does this cell span?
                span_el = tc_pr.find(f'{{{W}}}gridSpan')
                span = int(span_el.get(f'{{{W}}}val', '1')) if span_el is not None else 1
                
                # Cell width = sum of spanned column widths
                cell_w = sum(new_widths_twips[col_idx:col_idx+span]) if col_idx < len(new_widths_twips) else new_widths_twips[-1]
                
                tc_w = tc_pr.find(f'{{{W}}}tcW')
                if tc_w is not None:
                    tc_w.set(f'{{{W}}}w', str(cell_w))
                    tc_w.set(f'{{{W}}}type', 'dxa')
                
                col_idx += span
        
        patches += 1
        old_w_pt = [round(w/TWIP, 1) for w in current_widths]
        new_w_pt  = [round(w/TWIP, 1) for w in new_widths_twips]
        print(f"  [cols] Table {ti} (ncols={n_cols}): {old_w_pt} → {new_w_pt}")
    
    doc.save(out_path)
    print(f"  [cols] {patches} table(s) patched.")
    return patches


# ══════════════════════════════════════════════════════════════════════════════
# FIX 3 — Heading underline detection and injection
# ══════════════════════════════════════════════════════════════════════════════

def detect_underlined_text(pdf_doc, max_gap_pt=3.0):
    """
    Detect underlined text in the PDF.

    Primary method: PDF span flags (bit 2 = underline). Reliable for all fonts
    that set underline in the font descriptor.

    Secondary method (spatial): Finds horizontal drawing lines immediately below
    text spans. This catches underlines drawn as separate rect/line objects
    (e.g. PTCL offer letter headings). However, this method is NOISY for
    documents with heavy border drawing (e.g. salary slips with many table
    borders). It is automatically disabled per-page when the drawing count
    exceeds a threshold (≥50 drawings), since table borders cannot be reliably
    distinguished from underlines in that case.

    Known limitation: spatial detection may miss drawn underlines on pages with
    many border drawings. This is acceptable per the validation requirement:
    "If not inexpensive and reliable, document as a known limitation."

    Returns list of {page_no, text, y_bottom, x0, x1, source} for underlined spans.
    """
    underlined = []

    for pno in range(pdf_doc.page_count):
        page = pdf_doc[pno]

        drawings = page.get_drawings()
        many_drawings = len(drawings) >= 50  # likely a bordered-table document

        # Build horizontal lines only if spatial detection is feasible
        h_lines = []
        if not many_drawings:
            for d in drawings:
                rect = d.get("rect")
                if rect is not None:
                    rw = rect[2] - rect[0]
                    rh = abs(rect[3] - rect[1])
                    if rw > 10 and rh <= 2.0:
                        h_lines.append((rect[0], (rect[1]+rect[3])/2, rect[2]))
                for item in d.get("items", []):
                    if item[0] == "l":
                        p1, p2 = item[1], item[2]
                        if abs(p2.y - p1.y) < 1.5:
                            x_min, x_max = min(p1.x, p2.x), max(p1.x, p2.x)
                            if x_max - x_min > 10:
                                h_lines.append((x_min, (p1.y + p2.y)/2, x_max))

        # Get table bounding boxes for exclusion (only used in spatial mode)
        table_y_ranges = []
        if h_lines:
            for tbl in page.find_tables():
                table_y_ranges.append((tbl.bbox[1] - 2, tbl.bbox[3] + 2))

        def in_table(y0, y1):
            for ty0, ty1 in table_y_ranges:
                if y0 >= ty0 and y1 <= ty1:
                    return True
            return False

        text_dict = page.get_text("dict", flags=0)
        for b in text_dict["blocks"]:
            if b.get("type") != 0:
                continue
            for line in b.get("lines", []):
                for span in line.get("spans", []):
                    txt = span.get("text", "").strip()
                    if not txt or len(txt) < 3:
                        continue

                    font_size = span.get("size", 0)

                    # Primary: PDF underline flag (reliable)
                    flags = span.get("flags", 0)
                    if flags & 4:
                        underlined.append({
                            "page_no": pno,
                            "text": txt,
                            "y_baseline": span["bbox"][3],
                            "x0": span["bbox"][0],
                            "x1": span["bbox"][2],
                            "source": "pdf_flag",
                        })
                        continue

                    # Secondary: spatial detection (only when feasible)
                    if not h_lines:
                        continue
                    if font_size < 9:
                        continue
                    if in_table(span["bbox"][1], span["bbox"][3]):
                        continue

                    span_y_bottom = span["bbox"][3]
                    span_x0 = span["bbox"][0]
                    span_x1 = span["bbox"][2]
                    span_w  = span_x1 - span_x0

                    for lx0, ly, lx1 in h_lines:
                        gap = ly - span_y_bottom
                        if -1.0 <= gap <= max_gap_pt:
                            overlap = min(span_x1, lx1) - max(span_x0, lx0)
                            if overlap > span_w * 0.70 and (lx1 - lx0) > 20:
                                underlined.append({
                                    "page_no": pno,
                                    "text": txt,
                                    "y_baseline": span_y_bottom,
                                    "x0": span_x0,
                                    "x1": span_x1,
                                    "source": "spatial",
                                })
                                break

    return underlined



def patch_underlines(docx_path, pdf_doc, out_path):
    """
    Add underline formatting to DOCX runs that correspond to underlined text in the PDF.
    Match by text content — if the run text matches an underlined span, add w:u.
    """
    underlined_spans = detect_underlined_text(pdf_doc)
    
    # Build a set of underlined texts for O(1) lookup
    underlined_texts = set()
    for span in underlined_spans:
        underlined_texts.add(span["text"].strip())
    
    if not underlined_texts:
        print("  [underline] No underlined text detected in PDF.")
        shutil.copy2(docx_path, out_path)
        return 0
    
    print(f"  [underline] {len(underlined_texts)} unique underlined text(s): {list(underlined_texts)[:5]}")
    
    doc = Document(docx_path)
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    patches = 0
    
    for para in doc.element.body.iter(f'{{{W}}}p'):
        # Build full paragraph text
        para_text = ''.join(
            t.text or '' for t in para.findall(f'.//{{{W}}}t')
        ).strip()
        
        if not para_text:
            continue
        
        # Check if this paragraph's text matches an underlined span
        is_underlined = False
        for u_text in underlined_texts:
            if (para_text == u_text or
                u_text in para_text or
                para_text in u_text):
                is_underlined = True
                break
        
        if not is_underlined:
            continue
        
        # Add underline to all runs in this paragraph
        for run in para.findall(f'{{{W}}}r'):
            rPr = run.find(f'{{{W}}}rPr')
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                run.insert(0, rPr)
            
            # Only add if not already underlined
            existing_u = rPr.find(f'{{{W}}}u')
            if existing_u is None:
                u_el = OxmlElement('w:u')
                u_el.set(f'{{{W}}}val', 'single')
                rPr.append(u_el)
                patches += 1
    
    doc.save(out_path)
    print(f"  [underline] {patches} run(s) underlined.")
    return patches


# ══════════════════════════════════════════════════════════════════════════════
# FIX 4 — Table row height: exact → atLeast (prevents wrapped text from
#          overlapping the row/paragraph that follows it)
# ══════════════════════════════════════════════════════════════════════════════

def fix_table_row_overflow(docx_path, out_path):
    """
    pdf2docx gives every table row a fixed height (<w:trHeight w:hRule="exact">)
    that matches the tight single-line spacing measured from the source PDF.
    That's invisible as long as every cell's text still fits on one line in
    Word/LibreOffice/WPS -- but it doesn't always: small differences between
    the PDF's embedded font metrics and the font Word substitutes can make a
    cell that was one line in the PDF wrap to two lines in the DOCX. Because
    the row height is "exact", the row does NOT grow to fit that second
    line -- it stays pinned at the original single-line height, so the
    wrapped line renders UNDERNEATH whatever comes right after it (the next
    row, or the next paragraph). The text is still there in the file; it's
    just visually overlapped/garbled by what follows, which is exactly what
    can make a specific value look "missing" to someone looking at the
    rendered document (confirmed against the real salary-slip PDF: e.g. the
    "Domicile: NW - Khyber Pakhtunkhwa" and "Length of Service: ..." cells
    both wrap and overlap the row beneath them with the original exact
    heights pdf2docx assigns).

    Fix: switch hRule from "exact" to "atLeast" on every table row, in every
    table (including tables nested inside table cells). "atLeast" means
    "never shorter than this" -- for the overwhelming majority of rows,
    where the text already fits on one line, this is a no-op: the row
    renders at exactly the same height as before. Only a row whose content
    actually needs two lines will grow past that minimum instead of
    clipping/overlapping. This touches only the height-enforcement rule on
    existing rows -- no widths, text, or table-detection logic are touched,
    so it can't affect which rows exist or what they contain.
    """
    doc = Document(docx_path)
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    changed = 0
    for tr in doc.element.body.iter(f'{{{W}}}tr'):
        trPr = tr.find(qn('w:trPr'))
        if trPr is None:
            continue
        trHeight = trPr.find(qn('w:trHeight'))
        if trHeight is None:
            continue
        if trHeight.get(qn('w:hRule')) == 'exact':
            trHeight.set(qn('w:hRule'), 'atLeast')
            changed += 1

    doc.save(out_path)
    print(f"  [row_height] {changed} table row(s) switched exact→atLeast height.")
    return changed


# ══════════════════════════════════════════════════════════════════════════════
# FIX 5 — Non-integer OOXML measurement attributes (borders invisible in
#          stricter DOCX readers, e.g. mobile WPS Office)
# ══════════════════════════════════════════════════════════════════════════════

def fix_non_integer_measurements(docx_path, out_path):
    """
    Round every OOXML attribute that the spec defines as an integer
    measurement, but that pdf2docx wrote as a raw Python float, to a clean
    whole number.

    Confirmed by direct inspection of pdf2docx's own base output (before any
    of our other post-processing runs): border-size attributes come out
    like `<w:top w:sz="5.599999999999909" w:val="single" .../>` and width
    attributes like `<w:tcW w:w="546.00000000002"/>` -- i.e. pdf2docx
    multiplies a PDF point-based measurement (a float) by a unit factor
    (8 for border eighths-of-a-point, 20 for twips) and writes the raw
    result straight into the XML, with no rounding. Per OOXML
    (ECMA-376/ISO-29500), `w:sz` on a border is ST_EighthPointMeasure and
    `w:w` on widths/margins is ST_TwipsMeasure -- both integer types; a
    decimal value there is out-of-schema even though it happens to parse as
    a number.

    Word and LibreOffice tolerate this (they coerce to a number and ignore
    the fractional part) -- confirmed by converting this exact file with
    LibreOffice and seeing full, correctly drawn black grid lines around a
    table whose cells carry exactly this kind of malformed `w:sz`. A
    stricter reader is free to reject an attribute that fails schema
    validation instead, and silently drop just that border -- which would
    make a table that has real single-line borders in its XML render with
    *no* visible border at all, indistinguishable from plain aligned text.
    That matches a real-world report of a borderless-looking "missing"
    table that a LibreOffice-based render of the same file shows fully
    gridded (see server/TABLE_DETECTION_NOTES.md for the full writeup).

    This is a narrow, whitelisted pass: only `w:sz` (any element) and `w:w`
    (any element) attributes, plus `w:val` specifically on `w:trHeight`, are
    touched, and only when their value isn't already a clean integer.

    SECOND, SEPARATE FIX BUNDLED HERE (found while re-testing the border
    theory above against the real deployed server -- the float-sz fix alone
    did NOT make the table visible in WPS, which disproves that theory as
    the (sole) explanation and pointed at something else): every single
    `w:color` attribute pdf2docx writes -- all 273 of them in one real test
    file, covering every colored border AND every colored text run -- is
    written in CSS/HTML style, e.g. `w:color="#000000"`. OOXML's ST_HexColor
    type (ECMA-376/ISO-29500) is strictly 6 hex digits or the literal
    "auto" -- NO leading "#". `w:color="#000000"` is out-of-schema the same
    way the float `w:sz` values were. This is systemic (100% of colors in
    the file, not just table borders) and present from pdf2docx's very
    first base conversion, before any post-processing -- i.e. it's been in
    every version of this converter's output from the start, including the
    ones where the table-missing complaint was first raised. Word and
    LibreOffice tolerate the "#" (confirmed: LibreOffice renders these
    borders fine); a stricter reader is free to reject it. This is a
    stronger, more systemic candidate than the float-sz issue for why a
    table with genuinely correct border XML renders with no visible border
    in one specific app but not another -- and, being a plain string fix
    (strip a leading "#" from a 6-hex-digit value), it carries the same
    "can only help, never hurts" safety property.
    """
    doc = Document(docx_path)
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    SZ = f'{{{W}}}sz'
    WIDTH = f'{{{W}}}w'
    VAL = f'{{{W}}}val'
    COLOR = f'{{{W}}}color'
    TRHEIGHT = f'{{{W}}}trHeight'

    fixed = 0
    colors_fixed = 0

    def _round_attr(el, attr_qname):
        nonlocal fixed
        raw = el.get(attr_qname)
        if raw is None:
            return
        try:
            as_float = float(raw)
        except (TypeError, ValueError):
            return
        rounded = str(max(0, int(round(as_float))))
        if rounded != raw:
            el.set(attr_qname, rounded)
            fixed += 1

    def _fix_color(el):
        nonlocal colors_fixed
        raw = el.get(COLOR)
        if raw is None or not raw.startswith('#'):
            return
        stripped = raw[1:]
        if len(stripped) == 6 and all(c in '0123456789abcdefABCDEF' for c in stripped):
            el.set(COLOR, stripped)
            colors_fixed += 1

    for el in doc.element.body.iter():
        _round_attr(el, SZ)
        _round_attr(el, WIDTH)
        if el.tag == TRHEIGHT:
            _round_attr(el, VAL)
        _fix_color(el)

    doc.save(out_path)
    print(f"  [fix_measurements] {fixed} non-integer measurement attribute(s) rounded, "
          f"{colors_fixed} '#'-prefixed color value(s) fixed to valid OOXML hex.")
    return fixed + colors_fixed


# ══════════════════════════════════════════════════════════════════════════════
# FIX 7 — Shrink a table that pdf2docx sized past the page's printable width
#          (found via a logo rendering off the edge of the page)
# ══════════════════════════════════════════════════════════════════════════════

def fix_table_page_overflow(docx_path, out_path):
    """
    Root cause (found from a real report: "the logo is coming right at the
    page's edge, not like in the original PDF"): the header row's table
    (page title text in one cell, the district logo image in the other) has
    a left indent (`<w:tblInd w:w="1262"/>`, from pdf2docx's own base
    conversion -- positions the table starting at the title text's exact
    PDF x-offset) *and* a `<w:tblGrid>` sized to the page's full usable
    width (10754 twips) -- pdf2docx computes the two independently and
    never checks whether indent + width still fits on the page. Measured
    directly on the actual rendered PDF: left margin (558 twips) + this
    indent (1262) + the table's own width (10754) = 12574 twips = 628.7pt,
    against a 594.96pt-wide page -- the table's right edge lands 33.7pt
    past the page edge, dragging the (inline, right-aligned) logo in its
    right-hand cell along with it. Confirmed by rendering and measuring the
    logo's bounding box: it really does extend past the page boundary.

    This isn't limited to the one table with the logo -- any table that
    picked up both an indent and a full-page-width grid from pdf2docx would
    overflow the same way, image or not. So this is a general, final
    safety pass: for every top-level table, using the actual page size,
    margins, and this table's own indent (all read directly from the file,
    nothing assumed), compute how much width is actually available. If the
    table's <w:tblGrid> total is wider than that, scale every column width
    down by the same ratio so the table's right edge lands exactly at the
    page margin -- proportions between columns are preserved, only the
    table's overall size shrinks to fit. A table that already fits is left
    completely untouched.

    IMPORTANT: per-cell widths are scaled from each cell's OWN current
    <w:tcW> (old_w * ratio), never recomputed from <w:tblGrid> position +
    gridSpan. Found by testing: this step runs last and touches nearly
    every table, because pdf2docx gives almost every table a small nonzero
    <w:tblInd> (even 2-8 twips is enough to count as "overflow" by a
    couple twips). The first version reconstructed each cell's width as
    sum(scaled_gridCol[col_idx : col_idx+span]) -- correct for tables where
    every row shares one column layout, but for a table whose rows use
    DIFFERENT gridSpan patterns (e.g. a merged Gross Pay/Deductions/Net Pay
    summary row), <w:tblGrid> is not a real per-row model (see
    fix_merged_row_narrow_wraps's docstring) -- reconstructing from it threw
    away the row-local width adjustments that step had just made and
    silently reverted "-7,155.00" back to its wrapping width, even though
    fix_merged_row_narrow_wraps ran earlier in the same pipeline and its own
    log line reported the row as fixed. Scaling each cell's already-current
    width directly avoids ever needing to know how a row's cells line up
    with any other row's, so it can't re-collapse a row-local fix, and for
    ordinary tables (where tcW already equals gridCol at every position) it
    produces the exact same numbers as before.
    """
    doc = Document(docx_path)
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    def q(tag):
        return f'{{{W}}}{tag}'

    def _num(el, attr, default=0.0):
        if el is None:
            return default
        raw = el.get(attr)
        if raw is None:
            return default
        try:
            return float(raw)
        except ValueError:
            return default

    body = doc.element.body
    sectPr = body.find(q('sectPr'))
    if sectPr is None:
        doc.save(out_path)
        print("  [page_overflow] No sectPr found — nothing to check.")
        return 0

    pgSz = sectPr.find(q('pgSz'))
    pgMar = sectPr.find(q('pgMar'))
    page_w = _num(pgSz, q('w'), 11906)
    left_margin = _num(pgMar, q('left'), 1440)
    right_margin = _num(pgMar, q('right'), 1440)
    usable_w = page_w - left_margin - right_margin

    fixed_tables = 0
    for tbl in body.findall(q('tbl')):
        grid = tbl.find(q('tblGrid'))
        if grid is None:
            continue
        gridCols = grid.findall(q('gridCol'))
        if not gridCols:
            continue
        col_widths = [_num(gc, q('w')) for gc in gridCols]
        total = sum(col_widths)
        if total <= 0:
            continue

        tblPr = tbl.find(q('tblPr'))
        tblInd = tblPr.find(q('tblInd')) if tblPr is not None else None
        indent = _num(tblInd, q('w'), 0.0)

        available = usable_w - indent
        if available <= 0 or total <= available:
            continue  # fits (or indent alone already exceeds the page -- not this fix's job)

        ratio = available / total
        new_widths = [max(1, int(round(w * ratio))) for w in col_widths]
        # Rounding can leave the sum a few twips off target -- absorb that in the last column.
        diff = int(round(available)) - sum(new_widths)
        new_widths[-1] = max(1, new_widths[-1] + diff)

        for gc, w in zip(gridCols, new_widths):
            gc.set(q('w'), str(w))

        for tr in tbl.findall(q('tr')):
            for tc in tr.findall(q('tc')):
                tcPr = tc.find(q('tcPr'))
                if tcPr is None:
                    continue
                tcW = tcPr.find(q('tcW'))
                if tcW is None:
                    continue
                old_cell_w = _num(tcW, q('w'), 0.0)
                if old_cell_w <= 0:
                    continue
                new_cell_w = max(1, int(round(old_cell_w * ratio)))
                tcW.set(q('w'), str(new_cell_w))
                tcW.set(q('type'), 'dxa')

        fixed_tables += 1
        print(f"  [page_overflow] Table shrunk {int(total)}→{sum(new_widths)} twips "
              f"(indent={int(indent)}, usable page width={int(usable_w)}) to fit the page.")

    doc.save(out_path)
    print(f"  [page_overflow] {fixed_tables} table(s) resized to fit within the page margins.")
    return fixed_tables


# ══════════════════════════════════════════════════════════════════════════════
# FIX 6 — Widen table columns whose real text width doesn't fit, so numbers
#          don't wrap mid-value inside their own cell
# ══════════════════════════════════════════════════════════════════════════════

def fix_narrow_column_wraps(docx_path, out_path):
    """
    Root cause (confirmed on the real July-2026 salary slip, "Gross Pay /
    Deductions / Net Pay" row): pdf2docx positions each cell's text with a
    paragraph-level <w:ind w:left="..."/> that reproduces the text's exact
    horizontal offset inside its source PDF column. That indent eats into
    the cell's own width, but nothing in the pipeline (neither pdf2docx nor
    our own `cols` width patch) checks whether `cell_width - indent` still
    leaves enough room for the text itself.

    Measured directly from this file: the "-7,155.00 " cell is 1060 twips
    wide with a 244-twip left indent, leaving 816 twips for text --
    reportlab's real Times-Bold-10pt metrics say that exact string needs
    816.6 twips. A 0.6-twip (less than 1/30 of a point) shortfall is enough
    for Word/LibreOffice to wrap after the hyphen (a legal line-break
    point), splitting "-7,155.00" into "-" / "7,155.00" on two lines. The
    neighboring "105,752.00 " cell is the same story: 964 twips available
    after its indent, 950 needed -- 14 twips of slack, i.e. still inside
    normal font-substitution rounding error. Both wrapped in the actual
    LibreOffice rendering of this file.

    Fix: measure the real rendered width of every table cell's text with the
    real font/size it's set in (reportlab's built-in AFM metrics -- not a
    guess), add a small (1pt/20-twip) safety margin, and compare against
    `cell_width - indent`. Where a column is short, widen it by *borrowing*
    twips from OTHER columns in the same table that have measured spare
    room -- the table's total width, and therefore the overall page layout,
    never changes.

    IMPORTANT SAFETY RESTRICTION, found by testing this against a real
    11-column summary row (Gross Pay / Deductions / Net Pay, built from
    gridSpan-merged cells whose span boundaries differ from row to row):
    this table's own <w:tblGrid> lists 11 *equal* placeholder columns even
    though the actual rendered width of every cell comes from that cell's
    own <w:tcW>, which pdf2docx sets independently per cell and does not
    keep consistent with gridSpan sums when rows don't share the same span
    pattern. Reconciling width via the shared gridCol model there overwrote
    real, already-reasonable per-cell widths with nonsense derived from the
    meaningless uniform grid -- making cells *narrower*, the opposite of the
    intent, and a real regression caught only by re-rendering and looking.
    To make that class of bug structurally impossible rather than
    special-casing it: this fix only ever touches a table where every row
    has the exact same number of cells and the exact same gridSpan at every
    position (i.e. no two rows disagree about where column boundaries fall)
    -- there, and only there, gridCol really does mean one consistent thing
    per column across the whole table, so borrowing/lending width between
    columns is well-defined. Any table with row-to-row span differences is
    left completely untouched, exactly as it was before this fix existed --
    a known, non-data-loss cosmetic wrap left in place beats a "fix" that
    silently shrinks a column.
    """
    from reportlab.pdfbase.pdfmetrics import stringWidth

    doc = Document(docx_path)
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    def q(tag):
        return f'{{{W}}}{tag}'

    SAFETY_TWIPS = 20  # ~1pt cushion for font-substitution rounding

    def map_font(name, bold, italic):
        n = (name or '').lower()
        if 'times' in n or 'serif' in n:
            base = 'Times'
        elif 'courier' in n or 'mono' in n:
            base = 'Courier'
        else:
            base = 'Helvetica'
        if base == 'Times':
            if bold and italic:
                return 'Times-BoldItalic'
            if bold:
                return 'Times-Bold'
            if italic:
                return 'Times-Italic'
            return 'Times-Roman'
        suffix = ''
        if bold and italic:
            suffix = '-BoldOblique'
        elif bold:
            suffix = '-Bold'
        elif italic:
            suffix = '-Oblique'
        return base + suffix if suffix else base

    def cell_needed_twips(tc):
        """Widest single line's (text-width + paragraph indent) in twips, or 0 if empty/unmeasurable."""
        widest = 0
        for p in tc.findall(q('p')):
            pPr = p.find(q('pPr'))
            ind_left = ind_right = 0
            if pPr is not None:
                ind = pPr.find(q('ind'))
                if ind is not None:
                    for attr, var in ((q('left'), 'left'), (q('right'), 'right')):
                        raw = ind.get(attr)
                        if raw is not None:
                            try:
                                val = int(round(float(raw)))
                            except ValueError:
                                val = 0
                            if var == 'left':
                                ind_left = val
                            else:
                                ind_right = val
            line_pt = 0.0
            has_text = False
            for r in p.findall(q('r')):
                text = ''.join(t.text or '' for t in r.findall(q('t')))
                if not text:
                    continue
                has_text = True
                rPr = r.find(q('rPr'))
                font_name, bold, italic, size_pt = None, False, False, 10.0
                if rPr is not None:
                    rFonts = rPr.find(q('rFonts'))
                    if rFonts is not None:
                        font_name = rFonts.get(q('ascii'))
                    bold = rPr.find(q('b')) is not None
                    i_el = rPr.find(q('i'))
                    italic = i_el is not None and i_el.get(q('val')) != '0'
                    sz = rPr.find(q('sz'))
                    if sz is not None:
                        try:
                            size_pt = float(sz.get(q('val'), '20')) / 2.0
                        except ValueError:
                            pass
                font = map_font(font_name, bold, italic)
                try:
                    line_pt += stringWidth(text, font, size_pt)
                except Exception:
                    line_pt += len(text) * size_pt * 0.55
            if has_text:
                widest = max(widest, int(round(line_pt * 20)) + ind_left + ind_right)
        return widest + SAFETY_TWIPS if widest else 0

    body = doc.element.body
    fixed_tables = 0

    for tbl in body.findall(q('tbl')):
        grid = tbl.find(q('tblGrid'))
        if grid is None:
            continue
        gridCols = grid.findall(q('gridCol'))
        n_cols = len(gridCols)
        if n_cols == 0:
            continue

        rows = tbl.findall(q('tr'))

        # Safety gate: only touch tables with NO gridSpan anywhere, i.e. every
        # row has exactly n_cols cells, each occupying exactly one grid
        # column. That's the only shape where a table's <w:tblGrid> width is
        # guaranteed to equal every cell's own <w:tcW> at that position, so
        # "column N" means one unambiguous, consistent thing across every
        # row. (See the docstring: a table built from row-to-row-different
        # gridSpan merges does NOT have that guarantee -- its tblGrid can be
        # a meaningless uniform placeholder while real widths live on tcW --
        # and reconciling width against gridCol there corrupted real widths
        # in testing. Skipping those tables entirely is what makes that bug
        # impossible here, rather than something to remember to avoid.)
        has_any_span = False
        for tr in rows:
            tcs = tr.findall(q('tc'))
            if len(tcs) != n_cols:
                has_any_span = True
                break
            for tc in tcs:
                tcPr = tc.find(q('tcPr'))
                if tcPr is not None and tcPr.find(q('gridSpan')) is not None:
                    has_any_span = True
                    break
            if has_any_span:
                break
        if has_any_span:
            continue

        col_widths = [int(round(float(gc.get(q('w'), '0') or 0))) for gc in gridCols]
        col_needed = [0] * n_cols

        row_cell_info = []
        for tr in rows:
            infos = []
            for col_idx, tc in enumerate(tr.findall(q('tc'))):
                needed = cell_needed_twips(tc)
                infos.append((tc, col_idx, 1))
                col_needed[col_idx] = max(col_needed[col_idx], needed)
            row_cell_info.append(infos)

        deficits, slack = {}, {}
        for ci in range(n_cols):
            need = col_needed[ci]
            if not need:
                continue  # empty column everywhere -- no measurement, don't touch
            have = col_widths[ci]
            if need > have:
                deficits[ci] = need - have
            elif have > need:
                slack[ci] = have - need

        if not deficits or not slack:
            continue  # nothing to fix, or nowhere safe to borrow from

        total_deficit = sum(deficits.values())
        total_slack = sum(slack.values())
        borrow_ratio = min(1.0, total_slack / total_deficit)

        new_widths = list(col_widths)
        granted = 0
        for ci, deficit in deficits.items():
            grant = int(round(deficit * borrow_ratio))
            new_widths[ci] += grant
            granted += grant
        if granted <= 0:
            continue

        remaining = granted
        for ci, spare in sorted(slack.items(), key=lambda kv: -kv[1]):
            if remaining <= 0:
                break
            take = min(spare, remaining)
            new_widths[ci] -= take
            remaining -= take

        if new_widths == col_widths:
            continue

        for gc, w in zip(gridCols, new_widths):
            gc.set(q('w'), str(w))
        for infos in row_cell_info:
            for tc, col_idx, span in infos:
                tcPr = tc.find(q('tcPr'))
                if tcPr is None:
                    continue
                tcW = tcPr.find(q('tcW'))
                if tcW is None:
                    continue
                cell_w = sum(new_widths[col_idx:col_idx + span]) if col_idx < n_cols else new_widths[-1]
                tcW.set(q('w'), str(cell_w))
        fixed_tables += 1

    doc.save(out_path)
    print(f"  [narrow_cols] {fixed_tables} table(s) had column widths widened to stop mid-word wraps.")
    return fixed_tables


# ══════════════════════════════════════════════════════════════════════════════
# FIX 6b — Same mid-word-wrap problem, but for gridSpan-merged summary rows
#          that fix_narrow_column_wraps deliberately skips
# ══════════════════════════════════════════════════════════════════════════════

def fix_merged_row_narrow_wraps(docx_path, out_path):
    """
    `fix_narrow_column_wraps` (the `narrow_cols` step) refuses to touch any
    table where rows disagree about gridSpan, because that table's shared
    <w:tblGrid> is not a reliable "same column across every row" model --
    reconciling against it corrupted real widths in testing (see that
    function's docstring). That safety gate is correct, but it leaves a
    real, user-reported symptom unfixed: on the real July-2026 salary slip,
    the "Gross Pay / Deductions / Net Pay" summary table (row 0: Payable /
    Recovered / Exempted / Recoverable; row 1: Gross Pay / Deductions /
    Net Pay) has a DIFFERENT gridSpan pattern in each row, so the whole
    table was left untouched -- and "-7,155.00" and "105,752.00" kept
    wrapping mid-number ("-" / "7,155.00" and "105,752." / "00"), exactly
    as reported.

    FIRST ATTEMPT (kept here as a record because it failed instructively):
    edit only each cell's own <w:tcW>, borrowing slack row-locally so each
    row's own total stays constant, and leave <w:tblGrid> alone. That
    changed the file (verified: the saved XML had new tcW values) but
    produced ZERO visible difference on re-render -- confirmed by manually
    forcing +400 twips onto the exact deficient cells and re-rendering: the
    wrap was pixel-for-pixel identical. Root cause: this table's <w:tblPr>
    has <w:tblLayout w:type="fixed"/> with <w:tblW w:type="auto" w:w="0"/>.
    For a table shaped like this -- a "fixed" layout where different rows
    slice the same <w:tblGrid> into different gridSpan patterns -- both
    Word and LibreOffice render each cell's width as the SUM OF THE GRID
    COLUMNS ITS SPAN COVERS, not its own <w:tcW>. <w:tcW> is written by
    pdf2docx for every cell regardless, but it is decorative here: the
    actual column boundaries, in every row, come from <w:tblGrid> alone.
    So the real fix has to change <w:tblGrid>, not individual <w:tcW>.

    Real fix: for every row, walk its cells in order and derive which
    contiguous <w:tblGrid> columns each one occupies (col 0 for the first
    cell, spanning `gridSpan` columns; the next cell picks up where that
    one left off; etc.). This only works if every row's spans add up to
    exactly n_cols -- if any row doesn't, this table's structure isn't the
    simple "each row fully re-partitions the same n_cols" shape this fix
    assumes, and it's left untouched (same caution as the other step).
    For every cell in every row, measure its real needed text width (same
    reportlab AFM-metrics technique as `fix_narrow_column_wraps`, with the
    <w:b> "on unless w:val=0" bug fixed here -- the earlier code treated a
    present-but-disabled <w:b w:val="0"/> as bold, overestimating need).
    Spread that need evenly across the grid columns the cell covers, then
    take, for every grid column, the MAXIMUM of what any row's cell needs
    from it -- so satisfying row 1's "-7,155.00" can never shrink a column
    below what row 0 needs from that same column. If the resulting minimum
    total still leaves room under the table's current overall width, the
    leftover is handed back out proportionally to the columns' original
    widths, so the table's total width -- and hence the page layout --
    never changes, only how it's divided. Every cell's own <w:tcW> is then
    rewritten to match its row's new grid-derived width too, so the file's
    tcW and tblGrid stay mutually consistent (even though rendering was
    shown to only actually depend on tblGrid for these tables).

    Verified on the real file: row 1 needs (after fixing the bold bug)
    Gross Pay=1476, 112,907.00=970, Deductions=1614, -7,155.00=1081,
    Net Pay=1499, 105,752.00=1166, (empty)=0 twips against columns
    currently 1956/1956/1956/978/1956/978/973; row 0 needs Payable=878,
    23,037.74=870, Recovered till JUL-2026:=2214, 1,920.00=966,
    Exempted: 0.17-=1504, Recoverable=1192, 21,117.91=870 against
    978/1956/1956/1956/978/1956/973. Column-by-column max of both rows'
    per-grid-column share comfortably fits inside the table's actual total
    (10753 twips) with slack to spare, so this is feasible without
    shrinking the table.
    """
    from reportlab.pdfbase.pdfmetrics import stringWidth

    doc = Document(docx_path)
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    def q(tag):
        return f'{{{W}}}{tag}'

    SAFETY_TWIPS = 20

    def map_font(name, bold, italic):
        n = (name or '').lower()
        if 'times' in n or 'serif' in n:
            base = 'Times'
        elif 'courier' in n or 'mono' in n:
            base = 'Courier'
        else:
            base = 'Helvetica'
        if base == 'Times':
            if bold and italic:
                return 'Times-BoldItalic'
            if bold:
                return 'Times-Bold'
            if italic:
                return 'Times-Italic'
            return 'Times-Roman'
        suffix = ''
        if bold and italic:
            suffix = '-BoldOblique'
        elif bold:
            suffix = '-Bold'
        elif italic:
            suffix = '-Oblique'
        return base + suffix if suffix else base

    def is_on(el):
        """OOXML boolean toggle: element present AND val is not '0'/'false' means on."""
        if el is None:
            return False
        val = el.get(q('val'))
        return val not in ('0', 'false')

    def cell_needed_twips(tc):
        widest = 0
        for p in tc.findall(q('p')):
            pPr = p.find(q('pPr'))
            ind_left = ind_right = 0
            if pPr is not None:
                ind = pPr.find(q('ind'))
                if ind is not None:
                    for attr, var in ((q('left'), 'left'), (q('right'), 'right')):
                        raw = ind.get(attr)
                        if raw is not None:
                            try:
                                val = int(round(float(raw)))
                            except ValueError:
                                val = 0
                            if var == 'left':
                                ind_left = val
                            else:
                                ind_right = val
            line_pt = 0.0
            has_text = False
            for r in p.findall(q('r')):
                text = ''.join(t.text or '' for t in r.findall(q('t')))
                if not text:
                    continue
                has_text = True
                rPr = r.find(q('rPr'))
                font_name, bold, italic, size_pt = None, False, False, 10.0
                if rPr is not None:
                    rFonts = rPr.find(q('rFonts'))
                    if rFonts is not None:
                        font_name = rFonts.get(q('ascii'))
                    bold = is_on(rPr.find(q('b')))
                    italic = is_on(rPr.find(q('i')))
                    sz = rPr.find(q('sz'))
                    if sz is not None:
                        try:
                            size_pt = float(sz.get(q('val'), '20')) / 2.0
                        except ValueError:
                            pass
                font = map_font(font_name, bold, italic)
                try:
                    line_pt += stringWidth(text, font, size_pt)
                except Exception:
                    line_pt += len(text) * size_pt * 0.55
            if has_text:
                widest = max(widest, int(round(line_pt * 20)) + ind_left + ind_right)
        return widest + SAFETY_TWIPS if widest else 0

    def get_span(tc):
        tcPr = tc.find(q('tcPr'))
        if tcPr is None:
            return 1
        span_el = tcPr.find(q('gridSpan'))
        if span_el is None:
            return 1
        try:
            return max(1, int(span_el.get(q('val'), '1')))
        except ValueError:
            return 1

    body = doc.element.body
    fixed_tables = 0

    for tbl in body.findall(q('tbl')):
        grid = tbl.find(q('tblGrid'))
        if grid is None:
            continue
        gridCols = grid.findall(q('gridCol'))
        n_cols = len(gridCols)
        if n_cols == 0:
            continue

        rows = tbl.findall(q('tr'))

        # Only the complement of fix_narrow_column_wraps's safety gate:
        # a table must have SOME row-to-row gridSpan disagreement, or this
        # step has nothing to do here (the other step already handled it,
        # or there's nothing to fix either way).
        has_any_span = False
        for tr in rows:
            tcs = tr.findall(q('tc'))
            if len(tcs) != n_cols:
                has_any_span = True
                break
            for tc in tcs:
                if get_span(tc) != 1:
                    has_any_span = True
                    break
            if has_any_span:
                break
        if not has_any_span:
            continue

        # Derive each row's grid-column ranges. Only proceed if EVERY row's
        # cells, walked in order via their own gridSpan, land on exactly
        # n_cols total -- i.e. each row is a full, clean re-partition of
        # the same n_cols grid. If any row doesn't fit that shape, this
        # table isn't the pattern this fix targets; leave it untouched.
        row_ranges = []  # per row: list of (start, end, tc, needed)
        shape_ok = True
        for tr in rows:
            col = 0
            ranges = []
            for tc in tr.findall(q('tc')):
                span = get_span(tc)
                if col + span > n_cols:
                    shape_ok = False
                    break
                needed = cell_needed_twips(tc)
                ranges.append((col, col + span, tc, needed))
                col += span
            if not shape_ok or col != n_cols:
                shape_ok = False
                break
            row_ranges.append(ranges)
        if not shape_ok:
            continue

        # Per grid column, the max share any row's cell needs from it.
        col_min = [0.0] * n_cols
        for ranges in row_ranges:
            for start, end, _tc, needed in ranges:
                span = end - start
                if span <= 0 or not needed:
                    continue
                share = needed / span
                for i in range(start, end):
                    if share > col_min[i]:
                        col_min[i] = share

        orig_widths = [float(gc.get(q('w'), '0') or 0) for gc in gridCols]
        current_total = sum(orig_widths)
        total_min = sum(col_min)

        if current_total <= 0:
            continue

        if total_min > current_total:
            # Not enough room to satisfy every measured need without
            # growing the table -- best effort: scale every column's
            # minimum down proportionally rather than leave some cells
            # under-served in an unpredictable way, and skip claiming a
            # fix since wrapping may still occur.
            scale = current_total / total_min
            new_widths = [c * scale for c in col_min]
        else:
            slack = current_total - total_min
            orig_total = sum(orig_widths) or 1.0
            new_widths = [col_min[i] + slack * (orig_widths[i] / orig_total) for i in range(n_cols)]

        new_widths_i = [max(1, int(round(w))) for w in new_widths]
        diff = int(round(current_total)) - sum(new_widths_i)
        new_widths_i[-1] = max(1, new_widths_i[-1] + diff)

        if new_widths_i == [int(round(w)) for w in orig_widths]:
            continue  # already satisfies every constraint as-is

        for gc, w in zip(gridCols, new_widths_i):
            gc.set(q('w'), str(w))

        for ranges in row_ranges:
            for start, end, tc, _needed in ranges:
                tcPr = tc.find(q('tcPr'))
                if tcPr is None:
                    continue
                tcW = tcPr.find(q('tcW'))
                if tcW is None:
                    continue
                tcW.set(q('w'), str(sum(new_widths_i[start:end])))

        fixed_tables += 1

    doc.save(out_path)
    print(f"  [merged_row_wraps] {fixed_tables} table(s) had their grid columns "
          f"re-divided (per-row gridSpan-aware) to stop mid-word wraps in merged/summary rows.")
    return fixed_tables


# ══════════════════════════════════════════════════════════════════════════════
# Main pipeline
# ══════════════════════════════════════════════════════════════════════════════

def fix_section_breaks(docx_path, out_path):
    """
    Remove or convert spurious section breaks that pdf2docx inserts inside
    paragraph pPr elements. These appear when pdf2docx models multi-column
    layouts (e.g. a salary slip with two side-by-side wage tables as a
    2-column page section).

    Strategy:
    - Find all <w:sectPr> inside <w:pPr>.
    - If the type is 'nextPage' (or missing), change it to 'continuous'.
    - If the type is 'nextColumn', change it to 'continuous' as well.
    - Remove or reset the <w:cols> element to num=1 so multi-column layout
      is cancelled (the tables themselves handle the side-by-side appearance).
    - Leave the final document-level <w:sectPr> (child of <w:body>) untouched.
    """
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    with zipfile.ZipFile(docx_path, 'r') as z:
        names = z.namelist()
        files = {n: z.read(n) for n in names}

    doc_xml = etree.fromstring(files['word/document.xml'])
    body = doc_xml.find(f'{{{W}}}body')

    fixed = 0
    paras = body.findall(f'.//{{{W}}}p')
    for p in paras:
        pPr = p.find(f'{{{W}}}pPr')
        if pPr is None:
            continue
        sectPr = pPr.find(f'{{{W}}}sectPr')
        if sectPr is None:
            continue

        typeEl = sectPr.find(f'{{{W}}}type')
        current_type = typeEl.get(f'{{{W}}}val', '') if typeEl is not None else ''

        changed = False
        # Convert any page/column break to continuous
        if current_type in ('', 'nextPage', 'oddPage', 'evenPage', 'nextColumn'):
            if typeEl is None:
                typeEl = etree.Element(f'{{{W}}}type')
                sectPr.insert(0, typeEl)
            typeEl.set(f'{{{W}}}val', 'continuous')
            changed = True

        # Remove multi-column layout: reset cols to num=1
        colsEl = sectPr.find(f'{{{W}}}cols')
        if colsEl is not None:
            num = colsEl.get(f'{{{W}}}num', '1')
            if num != '1':
                # Set to single column
                colsEl.set(f'{{{W}}}num', '1')
                # Remove equalWidth and space if set
                for attr in list(colsEl.attrib.keys()):
                    if attr != f'{{{W}}}num':
                        del colsEl.attrib[attr]
                changed = True

        if changed:
            fixed += 1

    print(f"  [section_breaks] {fixed} section break(s) normalized.")

    files['word/document.xml'] = etree.tostring(
        doc_xml, xml_declaration=True, encoding='UTF-8', standalone=True)
    with zipfile.ZipFile(out_path, 'w', zipfile.ZIP_DEFLATED) as z:
        for name, data in files.items():
            z.writestr(name, data)

    return fixed


def postprocess(pdf_path, docx_in_path, docx_out_path, fixes=("section_breaks", "callout", "cols", "title_indent", "underline", "row_height", "narrow_cols", "merged_row_wraps", "fix_measurements", "page_overflow")):

    """
    Apply all enabled post-processing fixes to a pdf2docx-generated DOCX.
    Fixes are applied in sequence; each step reads from the previous output.
    """
    import tempfile
    
    print(f"\n{'='*64}")
    print(f"Post-processing: {os.path.basename(docx_in_path)}")
    print(f"  PDF:      {pdf_path}")
    print(f"  DOCX in:  {docx_in_path}")
    print(f"  DOCX out: {docx_out_path}")
    print(f"  Fixes:    {fixes}")
    print(f"{'='*64}")
    
    pdf_doc = pymupdf.open(pdf_path)
    
    # Use temp files for intermediate steps
    tmp_dir = tempfile.mkdtemp()
    step_in  = docx_in_path
    
    step_order = ["section_breaks", "callout", "cols", "title_indent", "underline", "row_height", "narrow_cols", "merged_row_wraps", "fix_measurements", "page_overflow"]
    applied_fixes = [f for f in step_order if f in fixes]
    
    for i, fix in enumerate(applied_fixes):
        is_last = (i == len(applied_fixes) - 1)
        step_out = docx_out_path if is_last else os.path.join(tmp_dir, f"step_{i}_{fix}.docx")
        
        if fix == "section_breaks":
            # Only normalize embedded section breaks for single-page source PDFs.
            # Multi-page PDFs use pPr sectPr legitimately to model page transitions.
            if pdf_doc.page_count == 1:
                n = fix_section_breaks(step_in, step_out)
            else:
                import shutil as _shutil
                _shutil.copy2(step_in, step_out)
                print(f"  [section_breaks] Skipped — source PDF has {pdf_doc.page_count} pages.")
                n = 0

        elif fix == "callout":
            n = inject_callout_boxes(step_in, pdf_doc, step_out)
        elif fix == "cols":
            n = patch_column_widths(step_in, pdf_doc, step_out)
        elif fix == "title_indent":
            n = fix_header_title_indent(step_in, step_out)
        elif fix == "underline":
            n = patch_underlines(step_in, pdf_doc, step_out)
        elif fix == "row_height":
            n = fix_table_row_overflow(step_in, step_out)
        elif fix == "narrow_cols":
            n = fix_narrow_column_wraps(step_in, step_out)
        elif fix == "merged_row_wraps":
            n = fix_merged_row_narrow_wraps(step_in, step_out)
        elif fix == "fix_measurements":
            n = fix_non_integer_measurements(step_in, step_out)
        elif fix == "page_overflow":
            n = fix_table_page_overflow(step_in, step_out)

        step_in = step_out
    
    # If no fixes applied, just copy
    if not applied_fixes:
        shutil.copy2(docx_in_path, docx_out_path)
    
    pdf_doc.close()
    shutil.rmtree(tmp_dir, ignore_errors=True)
    print(f"\n  Done → {docx_out_path}")


# ─── CLI ───────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    if len(sys.argv) < 4:
        print("Usage: python postprocess_docx.py <input.pdf> <input.docx> <output.docx> [fixes...]")
        print("  fixes: callout cols underline  (default: all three)")
        sys.exit(1)
    
    pdf_path  = sys.argv[1]
    docx_in   = sys.argv[2]
    docx_out  = sys.argv[3]
    fixes = tuple(sys.argv[4:]) if len(sys.argv) > 4 else ("callout", "cols", "title_indent", "underline", "row_height", "narrow_cols", "merged_row_wraps", "fix_measurements", "page_overflow")
    
    postprocess(pdf_path, docx_in, docx_out, fixes=fixes)
